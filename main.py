from scrape import extract_text, split_book_content
from parse import parse_with_ollama
import streamlit as st
from docx import Document
from io import BytesIO

st.title("Book Summarizer")
uploaded_file = st.file_uploader("Upload Desired Book:", type=["pdf"])
if uploaded_file is not None:
    st.write("Extracting text from the file...")
    st.session_state.book_text = extract_text(uploaded_file)
    st.session_state.book_chunks = split_book_content(st.session_state.book_text)
    st.write("Text extracted successfully!")

if "book_text" in st.session_state:
    if st.button("Parse Content"):
        st.write("Parsing the content")
        st.write("Number of lines: ", len(st.session_state.book_text.split('\n')))
        response = parse_with_ollama(st.session_state.book_text)
        st.write(response)
        
        # Create a docx file
        doc = Document()
        doc.add_heading('Parsed Content', 0)
        doc.add_paragraph(response)
        
        # Save the docx file to a BytesIO object
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Provide a download button
        st.download_button(
            label="Download Parsed Content as DOCX",
            data=buffer,
            file_name="parsed_content.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
